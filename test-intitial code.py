from docx import Document
import re
import pandas
from datetime import datetime
from main import date_mod, name


doc_2 = Document("test_initial.docx")

for i in range(len(date_mod)):
    if doc_2.tables[0].cell(9 + i, 0).text == doc_2.tables[0].cell(10 + i, 0).text:
        if doc_2.tables[0].cell(9 + i, 7).text == "":
            doc_2.tables[0].cell(9 + i, 7).text = doc_2.tables[0].cell(10 + i, 7).text

        elif doc_2.tables[0].cell(9 + i, 8).text == "":
            doc_2.tables[0].cell(9 + i, 8).text = doc_2.tables[0].cell(10 + i, 8).text

doc_2.save(f"test - {name}.docx")
