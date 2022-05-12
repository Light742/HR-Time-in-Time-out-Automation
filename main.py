from docx import Document
import re
import pandas
from datetime import datetime
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import time

Time_logs = pandas.read_csv("Timelogs.csv")

Time_in = Time_logs._Time.to_list()
f_key = Time_logs.Function_Key.to_list()

doc = Document("Test.docx")
name = "Jayson M. Hinggo"
group = "Alpha"
period_covered = "1-15"
employee_record = Time_logs[Time_logs.User_Name == name]
dates = employee_record._Time.to_list()

schedule = pandas.read_csv("Operations_schedule.csv")

if group == "Alpha":
    employee_schedule = schedule.Alpha
    employee_holiday = schedule.holiday_Alpha
elif group == "Bravo":
    employee_schedule = schedule.Bravo
    employee_holiday = schedule.holiday_Bravo
elif group == "Charlie":
    employee_schedule = schedule.Charlie
    employee_holiday = schedule.holiday_Charlie
elif group == "Delta":
    employee_schedule = schedule.Delta
    employee_holiday = schedule.holiday_Delta
else:
    print("group does not exist")
# how to extract time
time_all = []
for time in dates:
    date = re.findall(r"\d{1,2}:\d{1,2}:\d{2} \D{2}", time)
    for x in date:
        time_all.append(x)
time_mod = time_all[::-1]
date_all = []
for d in dates:
    date = re.findall(r"\d{1,2}/\d{1,2}/\d{4}", d)
    for x in date:
        date_all.append(x)
date_mod = date_all[::-1]

redundant_checker = []
for dates in date_mod:
    if dates not in redundant_checker:
        redundant_checker.append(dates)

User_ID = employee_record.User_ID.to_list()
employee_ID = str(User_ID[0])
fkey = employee_record.Function_Key.to_list()
fkey_final = fkey[::-1]
doc.tables[0].cell(0, 3).text = name
doc.tables[0].cell(1, 3).text = employee_ID
j = doc.tables[0].cell(7, 3)
k = doc.tables[0].cell(7, 11)
J = j.merge(k)

x = 0
for n in range(len(date_mod)):
    if 9 + n <= 9:
        doc.tables[0].cell(9 + n, 0).text = date_mod[n]
        day = datetime.strptime(date_mod[n], "%m/%d/%Y").strftime("%a").upper()
        doc.tables[0].cell(9 + n, 1).text = day
    else:
        doc.tables[0].add_row()
        new_row = doc.tables[0].rows[9 + x]._tr
        new_row.addnext(doc.tables[0].rows[-1]._tr)
        a = doc.tables[0].cell(10 + x, 2)
        b = doc.tables[0].cell(10 + x, 3)
        A = a.merge(b)
        c = doc.tables[0].cell(10 + x, 4)
        d = doc.tables[0].cell(10 + x, 5)
        C = c.merge(d)
        e = doc.tables[0].cell(10 + x, 6)
        f = doc.tables[0].cell(10 + x, 7)
        E = e.merge(f)
        g = doc.tables[0].cell(10 + x, 10)
        h = doc.tables[0].cell(10 + x, 11)
        G = g.merge(h)
        doc.tables[0].cell(9 + n, 0).text = date_mod[n]
        day = datetime.strptime(date_mod[n], "%m/%d/%Y").strftime("%a").upper()
        doc.tables[0].cell(9 + n, 1).text = day
        x += 1
F4_list = []
F1_list = []
out_list = []
for n in range(len(time_mod)):
    if fkey_final[n] == "F1":
        doc.tables[0].cell(9 + n, 7).text = time_mod[n]
        F1_list.append(time_mod[n])
    elif fkey_final[n] == "F4":
        doc.tables[0].cell(9 + n, 8).text = time_mod[n]
        F4_list.append(time_mod[n])

    else:
        doc.tables[0].cell(9 + n, 8).text = time_mod[n]
        out_list.append(time_mod[n])

F1_list.append("12:00:00 AM")
F1_list.append("12:00:00 PM")
out_list.append("12:00:00 AM")
out_list.append("12:00:00 PM")

for time in F4_list:
    out_list.append(time)

for time in F4_list:
    for i in range(len(doc.tables[0].rows) - 9):
        if doc.tables[0].cell(9 + i, 8).text == time:
            if doc.tables[0].cell(8 + i, 8) != "" and doc.tables[0].cell(8 + i, 7) != "":
                doc.tables[0].cell(8 + i, 8).text = "12:00:00 PM"
            doc.tables[0].cell(9 + i, 7).text = "12:00:00 AM"

for i in range(len(date_mod) - 1, -1, -1):
    if doc.tables[0].cell(9 + i, 0).text == doc.tables[0].cell(10 + i, 0).text:
        if doc.tables[0].cell(9 + i, 7).text == "":
            doc.tables[0].cell(9 + i, 7).text = doc.tables[0].cell(10 + i, 7).text
            table_element = doc.tables[0]._tbl
            table_element.remove(doc.tables[0].rows[10 + i]._tr)
        elif doc.tables[0].cell(9 + i, 8).text == "":
            doc.tables[0].cell(9 + i, 8).text = doc.tables[0].cell(10 + i, 8).text
            table_element = doc.tables[0]._tbl
            table_element.remove(doc.tables[0].rows[10 + i]._tr)


hours = []
for i in range(len(doc.tables[0].rows) - 10, -1, -1):
    if doc.tables[0].cell(9 + i, 7).text in F1_list:
        if doc.tables[0].cell(9 + i, 8).text in out_list:
            a = datetime.strptime(doc.tables[0].cell(9 + i, 7).text, "%I:%M:%f %p")
            b = datetime.strptime(doc.tables[0].cell(9 + i, 8).text, "%I:%M:%f %p")
            aa = float(a.strftime("%I")) + float(a.strftime("%M")) / 60
            if a.strftime("%p") == "PM":
                aa += 12
            if a.strftime("%I") == str(12) and a.strftime("%p") == "AM":
                aa -= 12

            bb = float(b.strftime("%I")) + float(b.strftime("%M")) / 60

            if b.strftime("%p") == "PM":
                bb += 12

            if b.strftime("%I") == str(12) and b.strftime("%p") == "AM":
                bb -= 12

            w_time = float(round((bb - aa), 2))

            zz_final = str(w_time)
            doc.tables[0].cell(9 + i, 9).text = zz_final
            print(zz_final)
    if period_covered == "1 - 15":
        track = doc.tables[0].cell(9 + i, 0).text
        text_test = re.search(r"15", track)
        if text_test is not None:
            table_element = doc.tables[0]._tbl
            table_element.remove(doc.tables[0].rows[9 + i]._tr)

final_dates = []
for i in range(len(doc.tables[0].rows) - 9):
    final_dates.append(doc.tables[0].cell(9 + i, 0).text)


for i in range(len(doc.tables[0].rows) - 9):
    hours.append(doc.tables[0].cell(9 + i, 9).text)

# payroll period
doc.tables[0].cell(7, 3).text = f"{date_mod[0]} - {date_mod[len(date_mod) - 1]}"

doc.tables[0].cell(8, 8).text = ""
doc.tables[0].cell(8, 8).paragraphs[0].add_run("TIME OUT").bold = True

if doc.tables[0].cell(9, 9).text == "24.0":
    table_element = doc.tables[0]._tbl
    table_element.remove(doc.tables[0].rows[9]._tr)

schedule_dates = schedule.Date.to_list()
for i in range(len(doc.tables[0].rows) - 9):
    if doc.tables[0].cell(9 + i, 0).text in schedule_dates:
        if employee_schedule[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Normal":
            if employee_holiday[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Regular Day":
                if hours[i] != "":
                    if float(hours[i]) > 8:
                        doc.tables[0].cell(9 + i, 10).text = str(round(float(hours[i]) - 8, 2)) + \
                                                             " hours overtime, Regular Day"
                    elif float(hours[i]) < 8:
                        if doc.tables[0].cell(9 + i, 8).text == "12:00:00 PM":
                            doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + " hours overtime, Regular Day"
                else:
                    doc.tables[0].cell(9 + i, 10).text = "insufficient data"
            else:
                doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + " hours overtime," + employee_holiday[
                    schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0]

        elif employee_schedule[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Day-Off":
            if employee_holiday[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Regular Day":
                if hours[i] != "":
                    doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + "hours overtime, Rest Day Duty "
                else:
                    doc.tables[0].cell(9 + i, 10).text = "insufficient data"
            else:
                if hours[i] != "":
                    doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + "hours overtime Rest Day Duty," + \
                                                         employee_holiday[schedule.Date ==
                                                                          doc.tables[0].cell(9 + i, 0).text].to_list()[
                                                             0]
                else:
                    doc.tables[0].cell(9 + i, 10).text = "insufficient data"
        elif employee_schedule[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Forced Overtime":
            if employee_holiday[schedule.Date == doc.tables[0].cell(9 + i, 0).text].to_list()[0] == "Regular Day":
                if hours[i] != "":
                    doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + "hours overtime, rest day duty"
                else:
                    doc.tables[0].cell(9 + i, 10).text = "insufficient data"
            else:
                if hours[i] != "":
                    doc.tables[0].cell(9 + i, 10).text = str(hours[i]) + "hours overtime," + \
                                                         employee_holiday[schedule.Date ==
                                                                          doc.tables[0].cell(9 + i, 0).text].to_list()[
                                                             0] + "rest day duty"
                else:
                    doc.tables[0].cell(9 + i, 10).text = "insufficient data"
signature = doc.tables[0].cell(len(doc.tables[0].rows) - 2, 8).paragraphs[2].add_run(f"{name}")
signature.font.size = Pt(10)
signature.font.bold = True
paragraph_format = doc.tables[0].cell(len(doc.tables[0].rows) - 2, 8).paragraphs[2].paragraph_format
paragraph_format.first_line_indent = Inches(-0.5)


for i in range(len(doc.tables[0].rows) - 9):
    if doc.tables[0].cell(9 + i, 0).text in date_mod:
        print("yes")
        for row in doc.tables[0].rows:
            row.height = Inches(0.05)





doc.save(f"test_{name}.docx")
